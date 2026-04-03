from pathlib import Path
from typing import Dict, List

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

from src.models import Article


def export_word(output_path: Path, articles: List[Article], summaries: Dict[str, str]) -> Path:
    """Export aggregation + summaries to Word document."""
    doc = Document()
    doc.add_heading("LLM-Assisted Cross-Platform Information Aggregation Report", level=1)

    doc.add_heading("1. Personalized Summary", level=2)
    for tag, text in summaries.items():
        doc.add_heading(f"Tag: {tag}", level=3)
        doc.add_paragraph(text)

    doc.add_heading("2. Aggregated Content", level=2)
    for idx, item in enumerate(articles, start=1):
        doc.add_heading(f"{idx}. {item.title}", level=3)
        doc.add_paragraph(f"Source: {item.source_type}/{item.source_name}")
        doc.add_paragraph(f"Link: {item.url}")
        doc.add_paragraph(f"Tags: {', '.join(item.tags)}")
        doc.add_paragraph(item.content[:500])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def export_pdf(output_path: Path, articles: List[Article], summaries: Dict[str, str]) -> Path:
    """Export report to PDF with Chinese font support."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c = canvas.Canvas(str(output_path), pagesize=A4)
    c.setFont("STSong-Light", 12)
    width, height = A4
    y = height - 40

    def write_line(line: str):
        nonlocal y
        if y < 40:
            c.showPage()
            c.setFont("STSong-Light", 12)
            y = height - 40
        c.drawString(40, y, line[:58])
        y -= 18

    write_line("LLM-Assisted Cross-Platform Information Aggregation Report")
    write_line("1. Personalized Summary")
    for tag, text in summaries.items():
        write_line(f"Tag: {tag}")
        for part in text.replace("\n", " ").split("."):
            if part.strip():
                write_line(part.strip() + ".")

    write_line("2. Aggregated Content")
    for idx, item in enumerate(articles[:120], start=1):
        write_line(f"{idx}. {item.title}")
        write_line(f"Source: {item.source_type}/{item.source_name}")
        write_line(f"Tags: {', '.join(item.tags)}")
        write_line(item.content[:120])

    c.save()
    return output_path
