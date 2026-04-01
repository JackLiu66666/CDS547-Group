import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Dict, List
from urllib.parse import unquote

import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


DATA_DIR = Path(__file__).parent / "data"
OUTPUT_DIR = Path(__file__).parent / "outputs"
TAG_FILE = DATA_DIR / "custom_tags.json"
SAMPLE_FILE = DATA_DIR / "sample_dataset.jsonl"


def normalize_text(text: str) -> str:
    text = re.sub(r"<[^>]+>", " ", text or "")
    text = re.sub(r"\s+", " ", text).strip()
    ad_keywords = ["广告", "推广", "点击领取", "扫码咨询", "赞助内容"]
    for kw in ad_keywords:
        text = text.replace(kw, "")
    return text.strip()


def extract_query_terms(query: str) -> List[str]:
    """Extract Chinese/English mixed query terms for relevance scoring."""
    q = (query or "").strip().lower()
    if not q:
        return []
    parts = re.findall(r"[\u4e00-\u9fa5A-Za-z0-9]{2,20}", q)
    if q not in parts:
        parts.append(q)
    unique = []
    for p in parts:
        if p not in unique:
            unique.append(p)
    return unique


def deduplicate_items(items: List[Dict]) -> List[Dict]:
    seen = set()
    result = []
    for item in items:
        item["title"] = normalize_text(item.get("title", ""))
        item["content"] = normalize_text(item.get("content", ""))
        key = item.get("url") or f"{item['title']}|{item['content'][:120]}"
        if not item["title"] or key in seen:
            continue
        seen.add(key)
        result.append(item)
    return result


def annotate_interest_tags(items: List[Dict], tags: List[str], user_interest: str) -> List[Dict]:
    """
    仅做兴趣标签标注，不影响搜索结果保留。
    """
    packed = tags + ([user_interest] if user_interest.strip() else [])
    packed = [x.strip() for x in packed if x.strip()]
    for item in items:
        text = f"{item.get('title', '')} {item.get('content', '')}".lower()
        matched = [tag for tag in packed if tag.lower() in text]
        item["interest_tags"] = matched if matched else ["未命中兴趣标签"]
    return items


def filter_by_search_keyword(items: List[Dict], keyword: str) -> List[Dict]:
    """
    仅按用户搜索词进行相关性筛选与排序，避免兴趣标签干扰搜索目标。
    """
    terms = extract_query_terms(keyword)
    if not terms:
        return items

    source_weight = {
        "GoogleNewsRSS": 1.0,
        "BingNewsRSS": 1.0,
        "arXiv": 1.15,
        "GitHub": 1.1,
        "HackerNews": 1.0,
        "Wikipedia": 0.95,
        "StackOverflow": 1.05,
    }
    scored = []
    for item in items:
        title = item.get("title", "").lower()
        content = item.get("content", "").lower()
        combined = f"{title} {content}"
        score = 0.0
        for term in terms:
            score += title.count(term) * 4.0
            score += content.count(term) * 1.5
            # URL含关键词通常表示高相关
            score += str(item.get("url", "")).lower().count(term) * 2.0

        # 标题前缀命中稍加权
        if any(title.startswith(t) for t in terms):
            score += 2.0

        # 长度惩罚，防止过长描述稀释
        if len(combined) > 600:
            score *= 0.9

        score *= source_weight.get(item.get("source_name", ""), 1.0)

        if score > 0:
            scored.append((score, item))

    if not scored:
        return items
    scored.sort(key=lambda x: x[0], reverse=True)
    return [x[1] for x in scored]


def filter_items_by_selected_keywords(items: List[Dict], selected_keywords: List[str]) -> List[Dict]:
    """
    用户根据推荐关键词做二次精准筛选。
    """
    picks = [k.strip().lower() for k in selected_keywords if k.strip()]
    if not picks:
        return items

    scored = []
    for item in items:
        text = f"{item.get('title', '')} {item.get('content', '')}".lower()
        score = 0.0
        for k in picks:
            score += text.count(k) * 2.0
            score += item.get("title", "").lower().count(k) * 2.5
        if score > 0:
            scored.append((score, item))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [x[1] for x in scored]


class TagManager:
    def __init__(self, file_path: Path = TAG_FILE):
        self.file_path = file_path
        self.file_path.parent.mkdir(parents=True, exist_ok=True)
        if not self.file_path.exists():
            self.save(["人工智能", "考研", "职场技能"])

    def load(self) -> List[str]:
        try:
            return json.loads(self.file_path.read_text(encoding="utf-8"))
        except Exception:
            return ["人工智能", "考研", "职场技能"]

    def save(self, tags: List[str]) -> None:
        clean = sorted(list(set([x.strip() for x in tags if x.strip()])))
        self.file_path.write_text(json.dumps(clean, ensure_ascii=False, indent=2), encoding="utf-8")

    def add(self, tag: str) -> List[str]:
        tags = self.load()
        if tag.strip() and tag.strip() not in tags:
            tags.append(tag.strip())
            self.save(tags)
        return self.load()

    def edit(self, old_tag: str, new_tag: str) -> List[str]:
        tags = [new_tag.strip() if x == old_tag else x for x in self.load()]
        self.save(tags)
        return self.load()

    def delete(self, tag: str) -> List[str]:
        tags = [x for x in self.load() if x != tag]
        self.save(tags)
        return self.load()


def ensure_dataset(size: int = 600) -> Path:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    if SAMPLE_FILE.exists():
        lines = [x for x in SAMPLE_FILE.read_text(encoding="utf-8").splitlines() if x.strip()]
        if len(lines) >= 500:
            return SAMPLE_FILE

    topics = [
        ("人工智能", "大模型", "知乎"),
        ("人工智能", "智能体", "学术摘要"),
        ("考研", "数学", "新闻"),
        ("考研", "英语", "公众号"),
        ("职场技能", "项目管理", "知乎"),
        ("职场技能", "数据分析", "新闻"),
    ]
    with SAMPLE_FILE.open("w", encoding="utf-8") as f:
        for i in range(size):
            tag, topic, source = topics[i % len(topics)]
            row = {
                "id": i + 1,
                "title": f"{tag}-{topic}信息第{i + 1}条",
                "content": f"这是一条关于{tag}和{topic}的演示数据，包含方法、案例、风险和执行建议。",
                "source_type": source,
                "source_name": "builtin",
                "publish_time": "2026-04-01",
                "url": f"https://example.com/{tag}/{i + 1}",
                "tags": [tag],
            }
            f.write(json.dumps(row, ensure_ascii=False) + "\n")
    return SAMPLE_FILE


def load_dataset(limit: int = 300) -> List[Dict]:
    ensure_dataset()
    data = []
    with SAMPLE_FILE.open("r", encoding="utf-8") as f:
        for i, line in enumerate(f):
            if i >= limit:
                break
            line = line.strip()
            if line:
                data.append(json.loads(line))
    return data


def summary_accuracy_estimate(items: List[Dict], summary: str) -> float:
    if not items or not summary.strip():
        return 0.0
    top_words = []
    for item in items[:50]:
        top_words.extend(re.findall(r"[\u4e00-\u9fa5A-Za-z0-9]{2,10}", item.get("title", "")))
    common = [w for w, _ in Counter(top_words).most_common(15) if len(w) >= 2]
    hit = sum(1 for w in common if w in summary)
    return round(min(1.0, hit / max(1, len(common))), 2)


def crawl_success_rate(stats: Dict[str, Dict]) -> float:
    if not stats:
        return 0.0
    total = len(stats)
    success = sum(1 for _, s in stats.items() if s.get("success"))
    return round(success / total, 2)


def to_dataframe(items: List[Dict]) -> pd.DataFrame:
    if not items:
        return pd.DataFrame()
    return pd.DataFrame(items)


def enrich_items_for_display(items: List[Dict]) -> List[Dict]:
    """Add display helpers like short source url and fallback tags."""
    enriched = []
    for item in items:
        clone = dict(item)
        raw_url = str(clone.get("url", ""))
        clone["display_url"] = unquote(raw_url[:120]) if raw_url else ""
        if "interest_tags" not in clone:
            clone["interest_tags"] = []
        enriched.append(clone)
    return enriched


def save_json(items: List[Dict], filename: str = "latest_results.json") -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    path.write_text(
        json.dumps(
            {
                "generated_at": datetime.now().isoformat(),
                "count": len(items),
                "items": items,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    return path


def export_word(items: List[Dict], summaries: Dict[str, str], filename: str = "report.docx") -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    doc = Document()
    doc.add_heading("跨平台信息聚合与个性化摘要报告", level=1)
    doc.add_heading("个性化摘要", level=2)
    for tag, text in summaries.items():
        doc.add_heading(tag, level=3)
        doc.add_paragraph(text)
    doc.add_heading("聚合结果", level=2)
    for idx, item in enumerate(items, start=1):
        doc.add_heading(f"{idx}. {item.get('title', '')}", level=3)
        doc.add_paragraph(f"来源: {item.get('source_type', '')} / {item.get('source_name', '')}")
        doc.add_paragraph(f"时间: {item.get('publish_time', '')}")
        doc.add_paragraph(f"链接: {item.get('url', '')}")
        tags = item.get("interest_tags", item.get("tags", []))
        doc.add_paragraph(f"标签: {', '.join(tags)}")
        doc.add_paragraph(item.get("content", "")[:500])
    doc.save(path)
    return path


def export_pdf(items: List[Dict], summaries: Dict[str, str], filename: str = "report.pdf") -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c = canvas.Canvas(str(path), pagesize=A4)
    c.setFont("STSong-Light", 11)
    _, h = A4
    y = h - 35

    def write(line: str):
        nonlocal y
        if y < 40:
            c.showPage()
            c.setFont("STSong-Light", 11)
            y = h - 35
        c.drawString(35, y, (line or "")[:60])
        y -= 16

    write("跨平台信息聚合与个性化摘要报告")
    write("个性化摘要")
    for tag, text in summaries.items():
        write(f"标签: {tag}")
        for part in text.replace("\n", " ").split("。"):
            if part.strip():
                write(part.strip() + "。")
    write("聚合结果")
    for idx, item in enumerate(items[:120], start=1):
        write(f"{idx}. {item.get('title', '')}")
        write(f"来源: {item.get('source_type', '')} / {item.get('source_name', '')}")
        tags = item.get("interest_tags", item.get("tags", []))
        if tags:
            write(f"标签: {', '.join(tags)}")
        write(item.get("content", "")[:100])
    c.save()
    return path
